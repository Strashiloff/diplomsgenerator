import docx
from docx.shared import Inches
from entity.diplom import Diplom
import os

images_dir = os.path.abspath('../images/')
docx_dir = os.path.abspath('../backend/MSword/')

def generateDocx(name, lastname, text):
    
    author = name + " " + lastname
    doc = docx.Document(docx_dir+'/template.docx')
    doc.paragraphs[6].text = text.topic 
    doc.paragraphs[15].text = author
    doc = ImportText(text, doc)
    namefile =  str(len(os.listdir(docx_dir)))+'out.docx'
    doc.save(docx_dir + '/'+namefile)
    return namefile

def ImportText(data , document):
    diplom = data
    indexphoto = 0
    listimg = os.listdir(images_dir)
    print(images_dir + listimg[0])
    for p in diplom.paragraphs:
        document.add_paragraph(p)
        document.add_picture(images_dir +'/'+ listimg[indexphoto], width=Inches(3))
        if len(listimg)>indexphoto+1: indexphoto += 1

    return document
    

#Пример ввода данных
#generateDocx('Вадим','Ляпушин', text(type Diplom))